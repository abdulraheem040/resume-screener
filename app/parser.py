"""
Resume parsing utilities.

Handles PDF and DOCX resumes, cleans the raw text, and splits it into
rough sections (skills / experience / education) using heading heuristics.
This is intentionally simple and regex-based -- good enough for a v1
pipeline. Swap in a proper NER-based resume parser later if you need
higher accuracy on things like job titles/dates.
"""

import re
import os
from dataclasses import dataclass, field


SECTION_HEADERS = {
    "skills": [
        "skills", "technical skills", "core competencies", "key skills",
    ],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment history", "work history",
    ],
    "education": [
        "education", "academic background", "qualifications",
    ],
    "summary": [
        "summary", "profile", "objective", "about me",
    ],
}


@dataclass
class ParsedResume:
    resume_id: str
    raw_text: str
    sections: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return self.raw_text


def _extract_pdf_text(file_path: str) -> str:
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def _extract_docx_text(file_path: str) -> str:
    import docx

    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]
    # Also pull text out of tables (many resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def clean_text(text: str) -> str:
    # Normalize whitespace, strip bullet characters, drop empty lines
    text = re.sub(r"[\u2022\u2023\u25E6\u2043\u2219•▪●○◦]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    lines = [line.strip() for line in text.split("\n")]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def split_sections(text: str) -> dict:
    """
    Very lightweight heading-based section splitter.
    Looks for lines that are short and match a known header keyword,
    then attributes subsequent lines to that section until the next
    header is found.
    """
    lines = text.split("\n")
    sections = {key: [] for key in SECTION_HEADERS}
    sections["other"] = []

    current_section = "other"
    for line in lines:
        lowered = line.lower().strip(":").strip()
        matched = None
        if len(lowered) <= 40:  # headers are short lines
            for section_name, keywords in SECTION_HEADERS.items():
                if any(lowered == kw or lowered.startswith(kw) for kw in keywords):
                    matched = section_name
                    break
        if matched:
            current_section = matched
            continue
        sections[current_section].append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items()}


def parse_resume_file(file_path: str, resume_id: str = None) -> ParsedResume:
    ext = os.path.splitext(file_path)[1].lower()
    resume_id = resume_id or os.path.basename(file_path)

    if ext == ".pdf":
        raw = _extract_pdf_text(file_path)
    elif ext in (".docx", ".doc"):
        raw = _extract_docx_text(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    cleaned = clean_text(raw)
    sections = split_sections(cleaned)
    return ParsedResume(resume_id=resume_id, raw_text=cleaned, sections=sections)


def parse_resume_bytes(data: bytes, filename: str, resume_id: str = None) -> ParsedResume:
    """Parse a resume from in-memory bytes (used by the FastAPI upload endpoint)."""
    import tempfile

    ext = os.path.splitext(filename)[1].lower()
    fd, tmp_path = tempfile.mkstemp(suffix=ext)
    with os.fdopen(fd, "wb") as f:
        f.write(data)
    try:
        return parse_resume_file(tmp_path, resume_id=resume_id or filename)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)